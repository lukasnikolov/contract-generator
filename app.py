from flask import Flask, request, send_file, jsonify
from docx import Document
import io
import tempfile
from flask_cors import CORS

app = Flask(__name__)
CORS(app)  # Povolit CORS pro API

@app.route('/')
def index():
    return "Contract generator is running."

@app.route('/api/generate', methods=['POST'])
def generate():
    data = request.json

    doc = Document("Rekapitulace_Ele.docx")
    placeholders = {
        "{{cislo_smlouvy}}": data.get("cislo_smlouvy", ""),
        "{{cislo_partnera}}": data.get("cislo_partnera", ""),
        "{{jmeno}}": data.get("jmeno", ""),
        "{{prijmeni}}": data.get("prijmeni", ""),
        "{{datum_narozeni}}": data.get("datum_narozeni", ""),
        "{{ulice_trvala}}": data.get("ulice_trvala", ""),
        "{{mesto_trvala}}": data.get("mesto_trvala", ""),
        "{{psc_trvala}}": data.get("psc_trvala", ""),
        "{{email}}": data.get("email", ""),
        "{{telefon}}": data.get("telefon", ""),
        "{{zpusob_odesilani}}": data.get("zpusob_odesilani", ""),
        "{{platby_faktury}}": data.get("platby_faktury", ""),
        "{{platby_zalohy}}": data.get("platby_zalohy", ""),
        "{{cislo_uctu}}": data.get("cislo_uctu", ""),
        "{{zahajeni_dodavek}}": data.get("zahajeni_dodavek", ""),
        "{{prolongace}}": data.get("prolongace", ""),
        "{{ean}}": data.get("ean", ""),
        "{{ulice_odber}}": data.get("ulice_odber", ""),
        "{{mesto_odber}}": data.get("mesto_odber", ""),
        "{{psc_odber}}": data.get("psc_odber", ""),
        "{{sazba}}": data.get("sazba", ""),
        "{{jistic}}": data.get("jistic", "")
    }

    for para in doc.paragraphs:
        for key, val in placeholders.items():
            if key in para.text:
                para.text = para.text.replace(key, val)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, val in placeholders.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, val)

    temp_docx = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp_docx.name)

    output = io.BytesIO()
    with open(temp_docx.name, "rb") as f:
        output.write(f.read())
    output.seek(0)

    return send_file(output, as_attachment=True, download_name="contract.docx")

if __name__ == "__main__":
    app.run(debug=False, host="0.0.0.0", port=10000)
